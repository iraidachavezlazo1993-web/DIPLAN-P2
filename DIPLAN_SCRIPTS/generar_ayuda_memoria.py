# -*- coding: utf-8 -*-
"""Genera un ayuda-memoria en Word con el contenido de las carpetas del proyecto."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SALIDA = os.path.join(ROOT, "DIPLAN_OUTPUT", "Ayuda_Memoria_DIPLAN.docx")

AZUL = RGBColor(0x1F, 0x3B, 0x73)
GRIS = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def titulo(texto, size=16, color=AZUL, space_before=10, space_after=6):
    p = doc.add_paragraph()
    p.space_before = Pt(space_before)
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def parrafo(texto, size=11, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def tabla(headers, filas):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for par in c.paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(10)
    for fila in filas:
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            celdas[i].text = str(v)
            for par in celdas[i].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9.5)
    return t


# ---------------- Portada / encabezado ----------------
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h.add_run("AYUDA MEMORIA")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = AZUL

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Procesamiento y consolidación de bases de intervenciones — DIPLAN")
r.font.size = Pt(12)
r.font.color.rgb = GRIS

parrafo("")
parrafo(
    "El presente documento describe la organización del proyecto en tres "
    "carpetas (insumos, scripts y resultados) y el contenido de cada una, con el "
    "fin de servir como guía rápida de referencia.", italic=True, color=GRIS)

# ---------------- Resumen de estructura ----------------
titulo("1. Estructura general", 15)
parrafo("El proyecto se organiza en tres carpetas principales:")
tabla(["Carpeta", "Contenido"], [
    ["DIPLAN_INPUTS", "Insumos de entrada: bases fuente y bases de cruce."],
    ["DIPLAN_SCRIPTS", "Programas en Python, en orden de ejecución (01, 02, 03)."],
    ["DIPLAN_OUTPUT", "Resultados generados: bases limpias, consolidadas, "
                      "base final, subcomponentes, diccionario y reporte."],
])

# ---------------- DIPLAN_INPUTS ----------------
titulo("2. Carpeta DIPLAN_INPUTS (insumos)", 15)
parrafo("Contiene 27 archivos. Se dividen en dos tipos:")

titulo("2.1. Bases fuente de los grupos", 12, color=GRIS, space_after=4)
tabla(["Archivo", "Grupo(s) que alimenta"], [
    ["ANIN.xlsx", "6 — Mantenimiento"],
    ["FONCODES.xlsx", "6 — Mantenimiento"],
    ["PEIP.xlsx / PEIP_solomant.xlsx", "3 — Módulos / 6 — Mantenimiento"],
    ["UGM.xlsx", "4 — Acondicionamiento, 6 — Mantenimiento"],
    ["UGME.xlsx", "1 — Mobiliario y equipo, 3 — Módulos"],
    ["UGRD.xlsx", "3 — Módulos"],
    ["UGSC.xlsx", "5 — Asistencias, 10 — Inspecciones"],
    ["2026.03.30 Zonales_UZ.xlsx", "5 — Asistencias, 10 — Inspecciones"],
    ["DIGEGED.xlsx", "6, 7 — Saneamiento, 8 — Prog. presup., 9 — Otros"],
    ["DRELM.xlsx", "1, 3, 6, 7, 9"],
    ["UE118.xlsx / UGEO.xlsx", "Bases de referencia adicionales"],
    ["Anexo 01 / Anexo 02 ... 300426", "1, 3, 4, 6, 11 — Demoliciones"],
    ["df_pi_listados_final_v7.xlsx", "2 — PI (base ya consolidada)"],
])

titulo("2.2. Bases de cruce y enriquecimiento", 12, color=GRIS, space_after=4)
tabla(["Archivo", "Función"], [
    ["Ordenamiento_ultimo.xlsx", "Define la composición de los 11 grupos."],
    ["df_vinculaciones_updated_20260527.xlsx", "Relaciona CUI / código modular con código local."],
    ["Copia de Padron_web.csv", "Imputa código local y datos geográficos."],
    ["ubigeo_UGEL.xlsx", "Departamento, provincia, distrito, DRE, UGEL, nombre IE."],
    ["ubigeo_pronied.xlsx", "Geografía, ruralidad y código modular."],
    ["ubigeo_cui.csv", "Geografía a partir del CUI."],
    ["Fichas_Indicadores_Cierre_Brecha_PNIE_060526.xlsx", "Catálogo de los 17 subcomponentes."],
    ["Rep_Activos_F7_02SET2024.zip", "Activo intervenido por CUI (subcomponentes)."],
    ["Rep_Inversiones_13ABR2026_EDU.xlsb", "Fecha de fin / cierre por CUI."],
    ["HR 079179-2026 - Inversiones.csv.gz", "Componentes de inversión por CUI."],
])

# ---------------- DIPLAN_SCRIPTS ----------------
titulo("3. Carpeta DIPLAN_SCRIPTS (programas)", 15)
parrafo("Tres programas que se ejecutan en orden:")
tabla(["Orden", "Script", "Qué hace"], [
    ["1", "DIPLAN_01_LIMPIEZA.py",
     "Limpia y normaliza cada base (fechas DD/MM/AAAA, montos numéricos, "
     "códigos estandarizados), las organiza en los 11 grupos, imputa el código "
     "local, agrega geografía (departamento/provincia/distrito/DRE/UGEL), "
     "asigna la fuente, uniformiza nombres de variables y elimina columnas "
     "duplicadas. Genera bases limpias, consolidadas por grupo, base final, "
     "diccionario y reporte."],
    ["2", "DIPLAN_02_ENRIQUECER.py",
     "Completa, en la base final, el tipo de activo y la fecha de fin "
     "(culminación / cierre) cruzando por CUI con los reportes de activos e "
     "inversiones."],
    ["3", "DIPLAN_03_SUBCOMPONENTES.py",
     "Evalúa, por código local, si se cumplen los 17 subcomponentes de cierre "
     "de brecha (1 = cumple, 0 = no), usando los activos y componentes de las "
     "bases de inversiones."],
])

# ---------------- DIPLAN_OUTPUT ----------------
titulo("4. Carpeta DIPLAN_OUTPUT (resultados)", 15)
parrafo("Todos los resultados se entregan en Excel (.xlsx) y en Parquet "
        "(.parquet, formato eficiente para grandes volúmenes).")

titulo("4.1. bases_limpias/", 12, color=GRIS, space_after=4)
parrafo("Una subcarpeta por cada grupo (Grupo_01 a Grupo_11). Dentro, cada base "
        "de origen ya limpia y normalizada.")

titulo("4.2. base_final/", 12, color=GRIS, space_after=4)
tabla(["Archivo / carpeta", "Contenido"], [
    ["base_final_armonizada", "Una sola tabla con columnas estándar "
     "(código local, CUI, modular, grupo, fuente, nombre IE, ubicación, tipo de "
     "intervención, tipo de activo, monto, devengado, fechas, etc.)."],
    ["base_final_union_completa", "Apila todas las bases conservando todas sus "
     "columnas originales (solo Parquet por su tamaño)."],
    ["consolidado_por_grupo/", "Una base consolidada por cada uno de los 11 "
     "grupos, con todos sus campos (Excel + Parquet)."],
    ["subcomponentes_total_matriz", "Una fila por código local con 1/0 para "
     "cada subcomponente, más los textos de activos y componentes."],
    ["subcomponentes_total_evidencia", "Detalle por código local y "
     "subcomponente: qué activo validó el cumplimiento y el estado de la inversión."],
    ["subcomponentes_total_diccionario", "Los 17 subcomponentes con sus "
     "palabras clave, fuente y regla de identificación."],
])

titulo("4.3. Archivos en la raíz de DIPLAN_OUTPUT", 12, color=GRIS, space_after=4)
tabla(["Archivo", "Contenido"], [
    ["diccionario_datos.xlsx", "Todas las variables de cada base, con su nombre "
     "estandarizado, nombre original, tipo de limpieza y un ejemplo."],
    ["reporte_limpieza.xlsx", "Resumen por base (filas, columnas, códigos "
     "imputados) y resumen por grupo."],
])

# ---------------- Nota final ----------------
titulo("5. Notas", 15)
parrafo("• Los 11 grupos corresponden a: 1 Mobiliario y equipo, 2 PI, "
        "3 Módulos, 4 Acondicionamiento, 5 Asistencias, 6 Mantenimiento, "
        "7 Saneamiento, 8 Programas presupuestales, 9 Otros, 10 Inspecciones, "
        "11 Demoliciones.")
parrafo("• La identificación de subcomponentes se realiza por coincidencia de "
        "palabras clave sobre los nombres de activos y componentes; por ello se "
        "incluyen los textos crudos, para permitir su revisión y mejora.")
parrafo("• Para reproducir todo el proceso, ejecutar los scripts en orden: "
        "01, luego 02 y finalmente 03.")

doc.save(SALIDA)
print("Generado:", SALIDA)
